"""
generate_docx_report.py
========================
Renders a ReportResult (from report_engine.py) into a polished, client-ready
Word document for the seller.

Uses python-docx (pure Python, pip install only - no Node/LibreOffice
dependency) so this runs on any machine an estate agent is likely to have.
"""

from __future__ import annotations

from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from report_engine import ReportResult

SEVERITY_COLORS = {
    "critical": RGBColor(0xC0, 0x1C, 0x1C),
    "warning": RGBColor(0xB4, 0x6A, 0x00),
    "positive": RGBColor(0x1E, 0x7A, 0x34),
    "info": RGBColor(0x33, 0x33, 0x33),
}
SEVERITY_LABELS = {
    "critical": "KEY ISSUE",
    "warning": "WORTH ADDRESSING",
    "positive": "WORKING WELL",
    "info": "FOR INFORMATION",
}


def _set_cell_background(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return h


def generate_report_docx(report: ReportResult, output_path: str,
                          agency_name: str = "", agent_contact: str = "") -> str:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Title page / header block ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Property Sale Review")
    run.bold = True
    run.font.size = Pt(26)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(report.listing.address)
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_text = f"Prepared {date.today().strftime('%d %B %Y')}"
    if agency_name:
        meta_text += f"  |  {agency_name}"
    if agent_contact:
        meta_text += f"  |  {agent_contact}"
    run = meta.add_run(meta_text)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # --- Summary box ---
    _add_heading(doc, "Summary", level=1)
    p = doc.add_paragraph(report.headline_summary)
    p.runs[0].font.size = Pt(12)

    # --- Key numbers table ---
    _add_heading(doc, "At a Glance", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].font.bold = True
        row[1].text = str(value)

    add_row("Asking price", f"£{report.listing.asking_price:,.0f}")
    if report.listing.original_asking_price != report.listing.asking_price:
        add_row("Original asking price", f"£{report.listing.original_asking_price:,.0f}")
    add_row("Days on market", report.days_on_market)
    if report.comparable_median:
        add_row("Local comparable median (sold price)", f"£{report.comparable_median:,.0f}")
        add_row("Comparables used", f"{report.comparable_count} sales in {report.listing.postcode.split()[0]}, last 36 months")
    if report.price_vs_median_pct is not None:
        sign = "+" if report.price_vs_median_pct >= 0 else ""
        add_row("Asking price vs. local median", f"{sign}{report.price_vs_median_pct}%")
    if report.area_trend and report.area_trend.get("latest"):
        lt = report.area_trend["latest"]
        add_row(f"{report.listing.region_for_trend} price trend (12mo)",
                f"{lt['annual_change_pct']:+.1f}% ({report.area_trend['trend_summary']})")
    if report.listing.listing_url:
        add_row("Listing", report.listing.listing_url)

    doc.add_paragraph()

    # --- Comparable sales table ---
    if report.comparables_sample:
        _add_heading(doc, "Recent Comparable Sales Nearby", level=1)
        doc.add_paragraph(
            "Sourced from HM Land Registry Price Paid Data (official, freely published sold "
            "prices for England & Wales)."
        ).runs[0].italic = True
        ctable = doc.add_table(rows=1, cols=4)
        ctable.style = "Light Grid Accent 1"
        hdr = ctable.rows[0].cells
        for i, label in enumerate(["Sold Date", "Price", "Type", "Address"]):
            hdr[i].text = label
            hdr[i].paragraphs[0].runs[0].font.bold = True
        for s in report.comparables_sample:
            row = ctable.add_row().cells
            row[0].text = s.date_sold
            row[1].text = f"£{s.price:,.0f}"
            row[2].text = s.property_type
            row[3].text = s.address
        doc.add_paragraph()

    # --- Findings ---
    _add_heading(doc, "Detailed Findings", level=1)
    order = {"critical": 0, "warning": 1, "info": 2, "positive": 3}
    for f in sorted(report.findings, key=lambda x: order.get(x.severity, 9)):
        p = doc.add_paragraph()
        label_run = p.add_run(f"[{SEVERITY_LABELS.get(f.severity, f.severity.upper())}] ")
        label_run.bold = True
        label_run.font.color.rgb = SEVERITY_COLORS.get(f.severity, RGBColor(0, 0, 0))
        headline_run = p.add_run(f.headline)
        headline_run.bold = True
        if f.detail:
            doc.add_paragraph(f.detail)

    # --- Action plan ---
    _add_heading(doc, "Recommended Action Plan", level=1)
    for i, step in enumerate(report.action_plan, start=1):
        doc.add_paragraph(f"{i}. {step}")

    # --- Footer note ---
    doc.add_paragraph()
    note = doc.add_paragraph()
    note_run = note.add_run(
        "This review uses official HM Land Registry sold-price and market trend data, "
        "combined with the current listing details, to assess likely reasons a property "
        "hasn't sold and what typically improves outcomes. It is a professional opinion to "
        "support a conversation with your estate agent, not a formal valuation."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.save(output_path)
    return output_path
